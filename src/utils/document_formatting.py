import re
import logging
import markdown2
from docx import Document
from html.parser import HTMLParser
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.styles.style import _ParagraphStyle
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# for removing the markdown and json from llm generated content
class DocumentParser:
    def __init__(self):
        self.style_manager = StyleManager()

    class DocxHTMLParser(HTMLParser):
        def __init__(self, doc, style_map):
            super().__init__()
            self.doc = doc
            self.style_map = style_map
            self.current_para = None
            self.current_style = "Normal"
            self.bold = False
            self.italic = False
            self.list_type = None
            self.list_level = 0

        def handle_starttag(self, tag, attrs):
            if tag.startswith("h") and tag[1:].isdigit():
                self.current_style = self.style_map.get(tag, "Normal")
                self.current_para = self.doc.add_paragraph()
                self.current_para.style = self.current_style
            elif tag == "p":
                self.current_style = "Normal"
                self.current_para = self.doc.add_paragraph()
            elif tag in ("ul", "ol"):
                self.list_level += 1
            elif tag == "li":
                self.current_para = self.doc.add_paragraph(style="List Bullet")
                self.current_para.paragraph_format.left_indent = Inches(
                    0.25 * self.list_level
                )
            elif tag in ("strong", "b"):
                self.bold = True
            elif tag in ("em", "i"):
                self.italic = True
            elif tag == "code":
                self.current_style = "Code"
                self.current_para = self.doc.add_paragraph()
                self.current_para.style = self.current_style

        def handle_endtag(self, tag):
            if tag in ("p", "li") or (tag.startswith("h") and tag[1:].isdigit()):
                self.current_para = None
            elif tag in ("ul", "ol"):
                self.list_level = max(0, self.list_level - 1)
            elif tag in ("strong", "b"):
                self.bold = False
            elif tag in ("em", "i"):
                self.italic = False

        def handle_data(self, data):
            if not data.strip():
                return
            if self.current_para is None:
                self.current_para = self.doc.add_paragraph()
            run = self.current_para.add_run(data)
            run.bold = self.bold
            run.italic = self.italic

    def _format_text(self, text: str, paragraph) -> None:
        try:
            doc = paragraph._parent
            html = markdown2.markdown(text)
            style_map = {
                "h1": "SectionHeader",
                "h2": "vignette",
                "h3": "Heading 4",
                "h4": "Heading 5",
                "h5": "Heading 6",
                "h6": "Heading 6",
            }
            parser = self.DocxHTMLParser(doc, style_map)
            parser.feed(html)
        except Exception:
            paragraph.add_run(text)

    def extract_json(self, text):
        match = re.search(r"```(?:json)?\\s*(\{.*?\})\\s*```", text, re.DOTALL)
        return match.group(1) if match else text.strip()


# for adding the styles to llm generated content
class StyleManager:
    def __init__(self):
        self.style_configs = {
            "CustomTitle": {
                "font_name": "Calibri Light",
                "size": 24,
                "bold": True,
                "alignment": WD_ALIGN_PARAGRAPH.CENTER,
                "space_after": 12,
            },
            "SectionHeader": {
                "font_name": "Calibri",
                "size": 16,
                "bold": True,
                "space_before": 12,
                "space_after": 6,
            },
            "Heading 3": {
                "font_name": "Calibri",
                "size": 14,
                "bold": True,
                "space_before": 12,
                "space_after": 4,
            },
            "Heading 4": {
                "font_name": "Calibri",
                "size": 12,
                "bold": True,
                "space_before": 10,
                "space_after": 4,
            },
            "Heading 5": {
                "font_name": "Calibri",
                "size": 11,
                "bold": True,
                "space_before": 8,
                "space_after": 4,
            },
            "Heading 6": {
                "font_name": "Calibri",
                "size": 11,
                "bold": True,
                "space_before": 6,
                "space_after": 4,
            },
            "Summary": {
                "font_name": "Calibri",
                "size": 11,
                "bold": False,
                "line_spacing": 1.15,
                "space_after": 6,
            },
            "List Bullet": {
                "font_name": "Calibri",
                "size": 11,
                "bold": False,
                "left_indent": 36,
                "space_after": 6,
            },
            "Code": {
                "font_name": "Courier New",
                "size": 10,
                "bold": False,
                "space_after": 6,
            },
        }

    def set_style_properties(self, style, font_name="Calibri", size=11, bold=False):
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold

    def init_document_styles(self, doc: Document) -> None:
        for style_name, config in self.style_configs.items():
            if style_name not in doc.styles:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            else:
                style = doc.styles[style_name]
            self.set_style_properties(
                style, config["font_name"], config["size"], config.get("bold", False)
            )
            if "alignment" in config:
                style.paragraph_format.alignment = config["alignment"]
            if "space_before" in config:
                style.paragraph_format.space_before = Pt(config["space_before"])
            if "space_after" in config:
                style.paragraph_format.space_after = Pt(config["space_after"])
            if "line_spacing" in config:
                style.paragraph_format.line_spacing = config["line_spacing"]
            if "left_indent" in config:
                style.paragraph_format.left_indent = Pt(config["left_indent"])

    def get_safe_style(self, doc: Document, style_name: str) -> _ParagraphStyle:
        return doc.styles.get(style_name, doc.styles["Normal"])


import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# This function converts folder updates to the input format required for merging documents
def convert_folder_updates_to_merge_input(
    folder_updates, output_filename, output_folder
):
    documents = []
    for folder_update in folder_updates:
        folder_name = folder_update["folder"]
        for doc in folder_update["documents"]:
            documents.append(
                {
                    "content": "",
                    "summary": doc.get("analysis", ""),
                    "blob_path": doc.get("blob_path", ""),
                    "file_name": doc.get("document_name", ""),
                    "folder_name": folder_name,
                }
            )
    return {
        "documents": documents,
        "output_filename": output_filename,
        "output_folder": output_folder,
    }
