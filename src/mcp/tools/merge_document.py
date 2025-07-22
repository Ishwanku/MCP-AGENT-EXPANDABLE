"""
Tool endpoint for merging summarized documents into a Word document.

This FastAPI route takes summarized documents, organizes them by folder, and generates a formatted Word (.docx) file as the final output.
"""
from docx import Document
from fastapi import APIRouter
from .models import (
    MergeDocumentRequest,
    MergeDocumentResponse,
)
from .utils import (
    DocumentParser,
    StyleManager,
    get_or_create_output_dir,
)
from datetime import datetime
from collections import defaultdict
from fastapi import HTTPException
import logging, string

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter()


# Tool to merge summarized documents into a Word document
@router.post("/merge_document", response_model=MergeDocumentResponse)
async def merge_document_tool(request: MergeDocumentRequest):
    logger.debug(f"Received merge document request: {request}")
    try:
        # Group documents by folder_name for section-wise organization
        grouped_documents = defaultdict(list)
        for doc in request.documents:
            grouped_documents[doc.folder_name].append(doc)

        # Create a new .docx document for the merged output
        doc = Document()
        style_manager = StyleManager()
        style_manager.init_document_styles(doc)
        parser = DocumentParser()
        section_letters = list(string.ascii_uppercase)

        # Iterate through grouped documents by folder
        for idx, (folder_name, docs_in_folder) in enumerate(grouped_documents.items()):
            section_label = (
                section_letters[idx]
                if idx < len(section_letters)
                else f"Section {idx+1}"
            )
            doc.add_heading(f"Section {section_label}: {folder_name}", level=1)

            for doc_info in docs_in_folder:
                doc.add_heading(f"Summary for {doc_info.file_name}:", level=2)
                paragraph = doc.add_paragraph()
                paragraph.style = "Summary"
                # Use summary if available, otherwise fallback to content
                text_to_format = doc_info.summary or doc_info.content
                if not text_to_format:
                    logger.warning(f"No summary or content for {doc_info.file_name}")
                    text_to_format = "No content or summary provided."
                parser._format_text(text_to_format, paragraph)

        # Save the merged document to the output folder
        output_dir = get_or_create_output_dir(request.output_folder)
        output_filename = (
            request.output_filename
            or f"summarized_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        output_path = output_dir / output_filename
        doc.save(output_path)

        # Return the path to the generated Word document
        return MergeDocumentResponse(output_file_path=str(output_path))
    except Exception as e:
        logger.error(f"Error in merge_document_tool: {str(e)}")
        # Raise an HTTPException if merging fails
        raise HTTPException(status_code=400, detail=f"Merge failed: {str(e)}")
