"""
Utility functions for document processing tools.

This module provides helpers for Azure Blob structure, document parsing and formatting, LLM summarization, parallel processing, and output directory management used by the document processing workflow.
"""
import re
import base64
import logging
import markdown2
from pathlib import Path
from docx import Document
from datetime import datetime
from typing import Dict, List
from langgraph.pregel import Send
from html.parser import HTMLParser
from docx.shared import Pt, Inches
from collections import defaultdict
from mcp.core.config import settings
from mcp.core.llm_client import LLMClient
from docx.enum.style import WD_STYLE_TYPE
from docx.styles.style import _ParagraphStyle
from docx.enum.text import WD_ALIGN_PARAGRAPH
from azure.identity import DefaultAzureCredential
from azure.storage.blob import BlobServiceClient, ContainerClient
from .models import DocumentState, SummarizeNodeResult, FolderUpdate, BatchState

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# This function helps to get folder and file structure from Azure Blob Storage to organize the generated summaries in folder(Section) wise structure
def get_blob_structure(storage_account, container_name):
    account_url = f"https://{storage_account}.blob.core.windows.net"
    blob_service_client = BlobServiceClient(
        account_url=account_url, credential=DefaultAzureCredential()
    )
    container_client: ContainerClient = blob_service_client.get_container_client(
        container_name
    )
    folder_map: Dict[str, List[str]] = {}
    blobs = container_client.list_blobs()
    for blob in blobs:
        blob_path = blob.name
        parts = blob_path.split("/")
        if len(parts) == 1:
            folder = "root"
            file_name = parts[0]
        else:
            folder = "/".join(parts[:-1])
            file_name = parts[-1]
        folder_map.setdefault(folder, []).append(file_name)
    return folder_map


# Define a safe base64 decode function to handle potential padding issues
def safe_base64_decode(s):
    missing_padding = len(s) % 4
    if missing_padding:
        s += "=" * (4 - missing_padding)
    return base64.b64decode(s).decode("utf-8")


# for removing the markdown and json from llm generated content
class DocumentParser:
    def __init__(self):
        self.style_manager = StyleManager()

    class DocxHTMLParser(HTMLParser):
        def __init__(self, doc, style_map):
            super().__init__()
            self.doc = doc
            self.style_map = style_map
            self.current_para = None
            self.current_style = "Normal"
            self.bold = False
            self.italic = False
            self.list_type = None
            self.list_level = 0

        def handle_starttag(self, tag, attrs):
            if tag.startswith("h") and tag[1:].isdigit():
                self.current_style = self.style_map.get(tag, "Normal")
                self.current_para = self.doc.add_paragraph()
                self.current_para.style = self.current_style
            elif tag == "p":
                self.current_style = "Normal"
                self.current_para = self.doc.add_paragraph()
            elif tag in ("ul", "ol"):
                self.list_level += 1
            elif tag == "li":
                self.current_para = self.doc.add_paragraph(style="List Bullet")
                self.current_para.paragraph_format.left_indent = Inches(
                    0.25 * self.list_level
                )
            elif tag in ("strong", "b"):
                self.bold = True
            elif tag in ("em", "i"):
                self.italic = True
            elif tag == "code":
                self.current_style = "Code"
                self.current_para = self.doc.add_paragraph()
                self.current_para.style = self.current_style

        def handle_endtag(self, tag):
            if tag in ("p", "li") or (tag.startswith("h") and tag[1:].isdigit()):
                self.current_para = None
            elif tag in ("ul", "ol"):
                self.list_level = max(0, self.list_level - 1)
            elif tag in ("strong", "b"):
                self.bold = False
            elif tag in ("em", "i"):
                self.italic = False

        def handle_data(self, data):
            if not data.strip():
                return
            if self.current_para is None:
                self.current_para = self.doc.add_paragraph()
            run = self.current_para.add_run(data)
            run.bold = self.bold
            run.italic = self.italic

    def _format_text(self, text: str, paragraph) -> None:
        try:
            doc = paragraph._parent
            html = markdown2.markdown(text)
            style_map = {
                "h1": "SectionHeader",
                "h2": "vignette",
                "h3": "Heading 4",
                "h4": "Heading 5",
                "h5": "Heading 6",
                "h6": "Heading 6",
            }
            parser = self.DocxHTMLParser(doc, style_map)
            parser.feed(html)
        except Exception:
            paragraph.add_run(text)

    def extract_json(self, text):
        match = re.search(r"```(?:json)?\\s*(\{.*?\})\\s*```", text, re.DOTALL)
        return match.group(1) if match else text.strip()


# for adding the styles to llm generated content
class StyleManager:
    def __init__(self):
        self.style_configs = {
            "CustomTitle": {
                "font_name": "Calibri Light",
                "size": 24,
                "bold": True,
                "alignment": WD_ALIGN_PARAGRAPH.CENTER,
                "space_after": 12,
            },
            "SectionHeader": {
                "font_name": "Calibri",
                "size": 16,
                "bold": True,
                "space_before": 12,
                "space_after": 6,
            },
            "Heading 3": {
                "font_name": "Calibri",
                "size": 14,
                "bold": True,
                "space_before": 12,
                "space_after": 4,
            },
            "Heading 4": {
                "font_name": "Calibri",
                "size": 12,
                "bold": True,
                "space_before": 10,
                "space_after": 4,
            },
            "Heading 5": {
                "font_name": "Calibri",
                "size": 11,
                "bold": True,
                "space_before": 8,
                "space_after": 4,
            },
            "Heading 6": {
                "font_name": "Calibri",
                "size": 11,
                "bold": True,
                "space_before": 6,
                "space_after": 4,
            },
            "Summary": {
                "font_name": "Calibri",
                "size": 11,
                "bold": False,
                "line_spacing": 1.15,
                "space_after": 6,
            },
            "List Bullet": {
                "font_name": "Calibri",
                "size": 11,
                "bold": False,
                "left_indent": 36,
                "space_after": 6,
            },
            "Code": {
                "font_name": "Courier New",
                "size": 10,
                "bold": False,
                "space_after": 6,
            },
        }

    def set_style_properties(self, style, font_name="Calibri", size=11, bold=False):
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold

    def init_document_styles(self, doc: Document) -> None:
        for style_name, config in self.style_configs.items():
            if style_name not in doc.styles:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            else:
                style = doc.styles[style_name]
            self.set_style_properties(
                style, config["font_name"], config["size"], config.get("bold", False)
            )
            if "alignment" in config:
                style.paragraph_format.alignment = config["alignment"]
            if "space_before" in config:
                style.paragraph_format.space_before = Pt(config["space_before"])
            if "space_after" in config:
                style.paragraph_format.space_after = Pt(config["space_after"])
            if "line_spacing" in config:
                style.paragraph_format.line_spacing = config["line_spacing"]
            if "left_indent" in config:
                style.paragraph_format.left_indent = Pt(config["left_indent"])

    def get_safe_style(self, doc: Document, style_name: str) -> _ParagraphStyle:
        return doc.styles.get(style_name, doc.styles["Normal"])


# This function is used to fan out documents for processing in parallel
async def fan_out_documents(state: BatchState):
    logger.info(f"Fan-out: Processing {len(state.documents)} documents")
    return [Send("summarize_node", {"documents": [doc]}) for doc in state.documents]


# This function summarizes each document using the LLM client
async def summarize_node(state: dict, max_tokens: int, temperature: float):
    logger.info(f"Summarize node called with state: {state}")
    doc = state["documents"][0]
    if isinstance(doc, dict):
        doc = DocumentState(**doc)
    logger.info(f"Summarizing document: {doc.file_name}")
    llm_client = LLMClient(settings)
    prompt = f"""
    Document: {doc.content}
    Please analyze the document and provide a concise summary.
    """
    try:
        analysis = await llm_client.generate_content(
            prompt=prompt, max_tokens=max_tokens, temperature=temperature
        )
        logger.info(f"LLM response for {doc.file_name}: {analysis}")
        if not analysis or analysis == "No response generated":
            analysis = "Summary generation failed."
            logger.warning(f"Summary generation failed for {doc.file_name}")
    except Exception as e:
        analysis = f"Summary generation failed: {str(e)}"
        logger.error(f"Error summarizing {doc.file_name}: {str(e)}")

    result = SummarizeNodeResult(
        folder_update=FolderUpdate(
            folder=doc.folder_name,
            documents=[
                {
                    "blob_path": doc.blob_path,
                    "document_name": doc.file_name,
                    "status": (
                        "summarized"
                        if analysis != "Summary generation failed."
                        else "failed"
                    ),
                    "analysis": analysis,
                }
            ],
        ),
        summary_text=f"Summary for {doc.file_name}:\n{analysis}",
    )
    logger.info(f"Summarize node result for {doc.file_name}: {result}")
    return {"summarize_node": result}


# This function aggregates results from the summarize_node and prepares the final output
def aggregate_results(state: BatchState):
    logger.info(f"Aggregating results from state: {state}")

    folder_updates_dict = defaultdict(list)
    summaries = []

    # Collect results from summarize_node
    for result in state.summarize_node:
        if hasattr(result, "folder_update"):
            folder = result.folder_update.folder
            folder_updates_dict[folder].extend(result.folder_update.documents)
            logger.info(f"Added documents to folder {folder}")
        if hasattr(result, "summary_text"):
            summaries.append(result.summary_text)
            logger.info(f"Added summary: {result.summary_text}")

    folder_updates = [
        FolderUpdate(folder=folder, documents=docs)
        for folder, docs in folder_updates_dict.items()
    ]

    logger.info(f"Aggregated folder updates: {len(folder_updates)} folders")

    result = {
        "folder_updates": folder_updates,
        "summaries": summaries,
        "documents": [],  # Clear to prevent further processing
        "summarize_node": state.summarize_node,  # Keep for inspection if needed
    }

    logger.info(f"Aggregation result: {result}")
    return result


# Create output directory
def get_or_create_output_dir(output_folder: str = None) -> Path:
    base_name = output_folder or getattr(settings, "OUTPUT_DIR", "output_folder")
    timestamp = datetime.now().strftime("on_%Y-%m-%d_at_%I_%M_%p")
    full_name = f"{base_name}_{timestamp}"
    output_dir = Path(full_name).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


# This function converts folder updates to the input format required for merging documents
def convert_folder_updates_to_merge_input(
    folder_updates, output_filename, output_folder
):
    documents = []
    for folder_update in folder_updates:
        folder_name = folder_update["folder"]
        for doc in folder_update["documents"]:
            documents.append(
                {
                    "content": "",
                    "summary": doc.get("analysis", ""),
                    "blob_path": doc.get("blob_path", ""),
                    "file_name": doc.get("document_name", ""),
                    "folder_name": folder_name,
                }
            )
    return {
        "documents": documents,
        "output_filename": output_filename,
        "output_folder": output_folder,
    }
